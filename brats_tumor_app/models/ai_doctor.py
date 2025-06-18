import os
from anthropic import Anthropic
from datetime import datetime
import json
from typing import Dict, List, Optional, Tuple
import pandas as pd
from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import tempfile
import numpy as np

class AIDoctor:
    def __init__(self, api_key: str):
        self.client = Anthropic(api_key=api_key)
        self.diagnosis_history = []
        self.medication_database = self._load_medication_database()
        
    def _load_medication_database(self) -> Dict:
        return {
            "感冒": ["对乙酰氨基酚", "布洛芬", "氨酚黄那敏", "感冒灵颗粒"],
            "发烧": ["对乙酰氨基酚", "布洛芬", "阿司匹林"],
            "咳嗽": ["右美沙芬", "氨溴索", "川贝枇杷膏"],
            "头痛": ["布洛芬", "对乙酰氨基酚", "阿司匹林"],
            "胃痛": ["奥美拉唑", "雷尼替丁", "铝碳酸镁"],
            "高血压": ["氨氯地平", "厄贝沙坦", "美托洛尔"],
            "糖尿病": ["二甲双胍", "格列美脲", "阿卡波糖"],
            "失眠": ["酒石酸唑吡坦", "地西泮", "褪黑素"],
            "过敏": ["氯雷他定", "西替利嗪", "地塞米松"],
            "腹泻": ["蒙脱石散", "诺氟沙星", "益生菌"],
        }
    
    def diagnose_stream(self, symptoms: str, medical_history: str = "", 
                     age: int = None, gender: str = None, 
                     vital_signs: Dict = None, lab_results: Dict = None):
        """流式诊断方法"""
        prompt = f"""你是一位专业的医生。请根据以下信息为患者进行诊断，以段落形式回复，重要信息用**加粗**显示：

患者信息：
- 年龄：{age if age else '未提供'}
- 性别：{gender if gender else '未提供'}
- 症状描述：{symptoms}
- 病史：{medical_history if medical_history else '无'}
- 生命体征：{json.dumps(vital_signs, ensure_ascii=False) if vital_signs else '未提供'}
- 检验结果：{json.dumps(lab_results, ensure_ascii=False) if lab_results else '未提供'}

请提供详细的医学分析，格式要求：
1. 用段落形式组织内容，避免过多的分点列表
2. 重要的诊断结果、疾病名称、治疗方案、注意事项等关键信息用**加粗**标记
3. 语言要专业但易懂，适合患者理解
4. 按以下顺序组织内容：初步诊断分析、可能病因、建议检查、治疗方案、重要注意事项

请以清晰、专业的段落形式回答，重要信息用**加粗**突出显示。"""

        try:
            with self.client.messages.stream(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7
            ) as stream:
                for text in stream.text_stream:
                    yield {
                        'type': 'content',
                        'content': text
                    }
                    
            # 获取完整响应用于历史记录
            final_message = stream.get_final_message()
            diagnosis_text = final_message.content[0].text
            
            # 保存诊断历史
            diagnosis_record = {
                'timestamp': datetime.now().isoformat(),
                'symptoms': symptoms,
                'diagnosis': diagnosis_text,
                'patient_info': {
                    'age': age,
                    'gender': gender,
                    'medical_history': medical_history
                }
            }
            self.diagnosis_history.append(diagnosis_record)
            
        except Exception as e:
            yield {
                'type': 'error',
                'content': f'诊断过程中出错: {str(e)}'
            }
    
    def diagnose(self, symptoms: str, medical_history: str = "", 
                 age: int = None, gender: str = None, 
                 vital_signs: Dict = None, lab_results: Dict = None) -> Dict:
        
        prompt = f"""你是一位专业的医生。请根据以下信息为患者进行诊断，以段落形式回复，重要信息用**加粗**显示：

患者信息：
- 年龄：{age if age else '未提供'}
- 性别：{gender if gender else '未提供'}
- 症状描述：{symptoms}
- 病史：{medical_history if medical_history else '无'}
- 生命体征：{json.dumps(vital_signs, ensure_ascii=False) if vital_signs else '未提供'}
- 检验结果：{json.dumps(lab_results, ensure_ascii=False) if lab_results else '未提供'}

请提供详细的医学分析，格式要求：
1. 用段落形式组织内容，不要使用大量的分点列表
2. 重要诊断结果、疾病名称、严重程度等关键信息用**加粗**标记
3. 语言要专业但易懂，适合患者理解

请以JSON格式返回，包含以下字段：
{{
    "preliminary_diagnosis": "**初步诊断结果**的段落描述，重要内容用**加粗**",
    "possible_diseases": [{{"disease": "疾病名", "probability": "概率"}}],
    "recommended_tests": ["检查项目的段落描述，重要检查用**加粗**"],
    "treatment_suggestions": ["治疗建议的段落描述，重要治疗方案用**加粗**"],
    "precautions": ["注意事项的段落描述，重要提醒用**加粗**"],
    "severity": "轻度/中度/重度",
    "referral_needed": true,
    "referral_department": "建议转诊科室"
}}"""

        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=2000,
                temperature=0.3,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # 获取响应文本
            response_text = response.content[0].text
            
            # 尝试提取JSON部分
            try:
                # 查找JSON的开始和结束
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = response_text[start_idx:end_idx]
                    diagnosis_result = json.loads(json_str)
                else:
                    raise ValueError("响应中未找到有效的JSON格式")
                    
            except json.JSONDecodeError:
                # 如果解析失败，尝试使用正则表达式清理响应
                import re
                # 移除可能的markdown代码块标记
                cleaned_text = re.sub(r'```json\s*', '', response_text)
                cleaned_text = re.sub(r'```\s*', '', cleaned_text)
                cleaned_text = cleaned_text.strip()
                
                # 再次尝试解析
                start_idx = cleaned_text.find('{')
                end_idx = cleaned_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = cleaned_text[start_idx:end_idx]
                    diagnosis_result = json.loads(json_str)
                else:
                    raise ValueError("无法从响应中提取有效的JSON")
            
            diagnosis_record = {
                "timestamp": datetime.now().isoformat(),
                "patient_info": {
                    "age": age,
                    "gender": gender,
                    "symptoms": symptoms,
                    "medical_history": medical_history,
                    "vital_signs": vital_signs,
                    "lab_results": lab_results
                },
                "diagnosis": diagnosis_result
            }
            
            self.diagnosis_history.append(diagnosis_record)
            
            return diagnosis_result
            
        except Exception as e:
            return {
                "error": f"诊断过程中出现错误：{str(e)}",
                "preliminary_diagnosis": "需要人工医生介入",
                "referral_needed": True
            }
    
    def prescribe_medication(self, diagnosis: str, patient_info: Dict) -> Dict:
        
        prompt = f"""作为专业医生，请根据以下诊断信息开具处方，回复要求以段落形式组织，重要信息用**加粗**显示：

诊断：{diagnosis}
患者信息：
- 年龄：{patient_info.get('age', '未知')}
- 性别：{patient_info.get('gender', '未知')}
- 过敏史：{patient_info.get('allergies', '无')}
- 当前用药：{patient_info.get('current_medications', '无')}

请提供处方建议，格式要求：
1. 用药说明和注意事项要用段落形式描述，不要大量使用分点列表
2. 重要的药物名称、剂量、警告信息用**加粗**标记
3. 语言要专业但患者易懂

必须严格按照以下JSON格式返回，不要包含任何其他文字：
{{
    "medications": [
        {{
            "name": "**药品名称**",
            "dosage": "**剂量**信息",
            "frequency": "**用药频率**",
            "duration": "**疗程**时长",
            "instructions": "用药说明的段落描述，重要注意点用**加粗**",
            "warnings": "注意事项的段落描述，重要警告用**加粗**"
        }}
    ],
    "drug_interactions": ["药物相互作用的段落描述，重要相互作用用**加粗**"],
    "follow_up": "复诊建议的段落描述，重要时间点和检查项目用**加粗**"
}}"""

        try:
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=1500,
                temperature=0.2,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # 获取响应文本
            response_text = response.content[0].text
            
            # 尝试提取JSON部分
            try:
                # 查找JSON的开始和结束
                start_idx = response_text.find('{')
                end_idx = response_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = response_text[start_idx:end_idx]
                    prescription = json.loads(json_str)
                else:
                    raise ValueError("响应中未找到有效的JSON格式")
                    
            except json.JSONDecodeError:
                # 如果解析失败，尝试使用正则表达式清理响应
                import re
                # 移除可能的markdown代码块标记
                cleaned_text = re.sub(r'```json\s*', '', response_text)
                cleaned_text = re.sub(r'```\s*', '', cleaned_text)
                cleaned_text = cleaned_text.strip()
                
                # 再次尝试解析
                start_idx = cleaned_text.find('{')
                end_idx = cleaned_text.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = cleaned_text[start_idx:end_idx]
                    prescription = json.loads(json_str)
                else:
                    raise ValueError("无法从响应中提取有效的JSON")
            
            # 验证必要的字段
            if 'medications' not in prescription:
                prescription['medications'] = []
            if 'follow_up' not in prescription:
                prescription['follow_up'] = '请按时复诊'
            if 'drug_interactions' not in prescription:
                prescription['drug_interactions'] = []
                
            return prescription
            
        except Exception as e:
            return {
                "error": f"开具处方时出现错误：{str(e)}",
                "medications": [],
                "drug_interactions": [],
                "follow_up": "请咨询人工医生"
            }
    
    
    def generate_medical_report(self, patient_info: Dict, diagnosis: Dict, 
                                prescription: Dict = None, format: str = "pdf") -> str:
        
        timestamp = datetime.now()
        report_id = f"MR_{timestamp.strftime('%Y%m%d_%H%M%S')}"
        
        if format == "docx":
            return self._generate_docx_report(report_id, patient_info, diagnosis, 
                                             prescription, timestamp)
        else:
            return self._generate_pdf_report(report_id, patient_info, diagnosis, 
                                           prescription, timestamp)
    
    def _generate_pdf_report(self, report_id: str, patient_info: Dict, 
                            diagnosis: Dict, prescription: Dict, 
                            timestamp: datetime) -> str:
        
        filename = f"{report_id}.pdf"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        doc = SimpleDocTemplate(filepath, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()
        
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            alignment=1
        )
        
        story.append(Paragraph("医疗诊断报告", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        info_data = [
            ["报告编号", report_id],
            ["生成时间", timestamp.strftime("%Y年%m月%d日 %H:%M")],
            ["患者姓名", patient_info.get('name', '---')],
            ["年龄", str(patient_info.get('age', '---'))],
            ["性别", patient_info.get('gender', '---')],
            ["联系电话", patient_info.get('phone', '---')]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.beige),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        
        story.append(info_table)
        story.append(Spacer(1, 0.3*inch))
        
        story.append(Paragraph("主诉症状", styles['Heading2']))
        story.append(Paragraph(patient_info.get('symptoms', '未提供'), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph("诊断结果", styles['Heading2']))
        story.append(Paragraph(f"初步诊断：{diagnosis.get('preliminary_diagnosis', '---')}", styles['Normal']))
        story.append(Paragraph(f"严重程度：{diagnosis.get('severity', '---')}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        if diagnosis.get('possible_diseases'):
            story.append(Paragraph("可能的疾病：", styles['Normal']))
            for disease in diagnosis.get('possible_diseases', []):
                story.append(Paragraph(f"  • {disease['disease']} (可能性：{disease['probability']})", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        if prescription:
            story.append(Paragraph("处方信息", styles['Heading2']))
            for med in prescription.get('medications', []):
                story.append(Paragraph(f"药品：{med['name']}", styles['Normal']))
                story.append(Paragraph(f"  剂量：{med['dosage']}", styles['Normal']))
                story.append(Paragraph(f"  频率：{med['frequency']}", styles['Normal']))
                story.append(Paragraph(f"  疗程：{med['duration']}", styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        
        story.append(Paragraph("医嘱和建议", styles['Heading2']))
        for suggestion in diagnosis.get('treatment_suggestions', []):
            story.append(Paragraph(f"  • {suggestion}", styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        story.append(Paragraph("注意事项", styles['Heading2']))
        for precaution in diagnosis.get('precautions', []):
            story.append(Paragraph(f"  • {precaution}", styles['Normal']))
        
        story.append(Spacer(1, 0.5*inch))
        story.append(Paragraph("医生签名：_________________", styles['Normal']))
        story.append(Paragraph("日期：_________________", styles['Normal']))
        
        doc.build(story)
        return filepath
    
    def _generate_docx_report(self, report_id: str, patient_info: Dict, 
                             diagnosis: Dict, prescription: Dict, 
                             timestamp: datetime) -> str:
        
        filename = f"{report_id}.docx"
        filepath = os.path.join(tempfile.gettempdir(), filename)
        
        doc = Document()
        
        title = doc.add_heading('医疗诊断报告', 0)
        title.alignment = 1
        
        doc.add_heading('基本信息', level=1)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid'
        
        info_items = [
            ("报告编号", report_id),
            ("生成时间", timestamp.strftime("%Y年%m月%d日 %H:%M")),
            ("患者姓名", patient_info.get('name', '---')),
            ("年龄", str(patient_info.get('age', '---'))),
            ("性别", patient_info.get('gender', '---')),
            ("联系电话", patient_info.get('phone', '---'))
        ]
        
        for i, (label, value) in enumerate(info_items):
            table.cell(i, 0).text = label
            table.cell(i, 1).text = value
        
        doc.add_heading('主诉症状', level=1)
        doc.add_paragraph(patient_info.get('symptoms', '未提供'))
        
        doc.add_heading('诊断结果', level=1)
        doc.add_paragraph(f"初步诊断：{diagnosis.get('preliminary_diagnosis', '---')}")
        doc.add_paragraph(f"严重程度：{diagnosis.get('severity', '---')}")
        
        if diagnosis.get('possible_diseases'):
            doc.add_paragraph("可能的疾病：")
            for disease in diagnosis.get('possible_diseases', []):
                doc.add_paragraph(f"  • {disease['disease']} (可能性：{disease['probability']})")
        
        if prescription:
            doc.add_heading('处方信息', level=1)
            for med in prescription.get('medications', []):
                doc.add_paragraph(f"药品：{med['name']}")
                doc.add_paragraph(f"  剂量：{med['dosage']}")
                doc.add_paragraph(f"  频率：{med['frequency']}")
                doc.add_paragraph(f"  疗程：{med['duration']}")
                doc.add_paragraph("")
        
        
        doc.add_heading('医嘱和建议', level=1)
        for suggestion in diagnosis.get('treatment_suggestions', []):
            doc.add_paragraph(f"  • {suggestion}")
        
        doc.add_heading('注意事项', level=1)
        for precaution in diagnosis.get('precautions', []):
            doc.add_paragraph(f"  • {precaution}")
        
        doc.add_paragraph("")
        doc.add_paragraph("医生签名：_________________")
        doc.add_paragraph("日期：_________________")
        
        doc.save(filepath)
        return filepath
    
    def get_diagnosis_history(self) -> List[Dict]:
        return self.diagnosis_history